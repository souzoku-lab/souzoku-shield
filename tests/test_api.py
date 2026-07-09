from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

from docx import Document
from fastapi.testclient import TestClient

from app import agent_run
from app.main import app


def test_health_and_case_payload() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        health = client.get("/api/health")
        assert health.status_code == 200
        assert health.json()["ok"] is True
        assert health.json()["llm_required"] is False

        case = client.get("/api/case")
        assert case.status_code == 200
        body = case.json()
        assert body["analysis"]["acquirer"]["id"] == "co_resident"
        assert body["analysis"]["case"]["land"]["estimated_value_yen"] == 70000000
        assert body["analysis"]["draft"]["section_5_overall_opinion"] == ""
        assert body["harness"]["ok"] is True
        assert body["harness"]["impact_note"] == "課税価格 +5,600万円（失われる評価減）"


def test_case_patch_changes_counterfactual_branch() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.patch("/api/case", json={"acquirer_type": "spouse"})
        assert response.status_code == 200
        body = response.json()
        assert body["analysis"]["acquirer"]["id"] == "spouse"
        required = body["analysis"]["acquirer"]["required_document_ids"]
        assert "co_resident_continuation" not in required

        branches = body["counterfactuals"]
        assert any(branch["to"] == "house_lost" for branch in branches)


def test_document_patch_moves_meter() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        before = client.get("/api/case").json()["analysis"]["completion"]["percent"]
        response = client.patch("/api/documents/co_resident_continuation", json={"status": "verified"})
        assert response.status_code == 200
        after = response.json()["analysis"]["completion"]["percent"]
        assert after > before


def test_static_ui_is_served() -> None:
    with TestClient(app) as client:
        index = client.get("/")
        script = client.get("/static/app.js")
        assert index.status_code == 200
        assert "小規模宅地 要件確認＆書面添付資料作成" in index.text
        assert "Review作成" in index.text
        assert "カード内容でReview作成" in index.text
        assert "相談文を実行" in index.text
        assert "① カード登録・Review作成" in index.text
        assert "② レビュー完了（承認）" in index.text
        assert "③ 書面添付資料をWord出力" in index.text
        assert "approvalHint" in index.text
        assert "wordHint" in index.text
        assert "wordFlowStatus" in index.text
        assert "wordFlowHint" in index.text
        assert "cardReviewButton" in index.text
        assert "inlineApproveButton" in index.text
        assert "inlineWordLink" in index.text
        assert "否認インパクトハーネス" in index.text
        assert "heirBoard" in index.text
        assert "heirRegistrationForm" in index.text
        assert "関係性" in index.text
        assert "同居・非同居" in index.text
        option_values = [
            segment.split('value="', 1)[1].split('"', 1)[0]
            for segment in index.text.split("<option ")[1:8]
        ]
        assert option_values == [
            "spouse",
            "eldest_son",
            "eldest_daughter",
            "second_son",
            "second_daughter",
            "third_son",
            "third_daughter",
        ]
        assert 'value="mother"' not in index.text
        assert "runError" in index.text
        assert index.text.index("heirBoard") < index.text.index("consultationText")
        assert "書面添付資料出力（Word）" in index.text
        assert "eligibilityImpactValue" in index.text
        assert "harnessImpactLabel" in index.text
        assert "harnessImpactValue" in index.text
        assert "badFixture" not in index.text
        assert "誤実装サンプル" not in index.text
        assert script.status_code == 200
        assert "現在案件インパクト" in script.text
        assert "課税価格" in script.text
        assert "Gemini 3.5 Flash 実接続" in script.text
        assert "カード内容でReview作成" in script.text
        assert "レビュー完了（承認）" in script.text
        assert "レビュー完了するとWord出力できます" in script.text
        assert "③ Word出力できます" in script.text
        assert "書面添付資料のdocxがダウンロード" in script.text
        assert "addHeir" in script.text
        assert "runCardReview" in script.text
        assert "誤実装サンプル" not in script.text


def test_heir_selection_drives_acquirer_and_ineligibility_alert() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        selected = client.patch("/api/case", json={"home_acquirer_id": "second_son"})
        relaxed = client.patch("/api/heirs/eldest_son", json={"co_resident": False})

    assert selected.status_code == 200
    selected_body = selected.json()
    assert selected_body["state"]["home_acquirer_id"] == "second_son"
    assert selected_body["analysis"]["home_acquirer"]["name"] == "次男"
    assert selected_body["analysis"]["acquirer"]["id"] == "house_lost"
    assert len(selected_body["analysis"]["eligibility_alerts"]) == 1
    assert selected_body["analysis"]["eligibility_alerts"][0]["id"] == "co_resident_and_spouse_block_house_lost_acquirer"
    assert "同居親族（長男）" in selected_body["analysis"]["eligibility_alerts"][0]["message"]
    assert "被相続人に配偶者（母）" in selected_body["analysis"]["eligibility_alerts"][0]["message"]
    assert selected_body["analysis"]["eligibility_alerts"][0]["impact_yen"] == 56000000
    assert any(
        "適用不可アラート" in line
        for line in selected_body["analysis"]["draft"]["section_3_land_review"]
    )

    assert relaxed.status_code == 200
    relaxed_body = relaxed.json()
    assert relaxed_body["analysis"]["home_acquirer"]["name"] == "次男"
    assert relaxed_body["analysis"]["acquirer"]["id"] == "house_lost"
    assert relaxed_body["analysis"]["eligibility_alerts"][0]["id"] == "spouse_blocks_house_lost_acquirer"
    assert "被相続人に配偶者（母）" in relaxed_body["analysis"]["eligibility_alerts"][0]["message"]


def test_selected_heir_co_residence_toggle_recalculates_acquirer() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.patch("/api/case", json={"home_acquirer_id": "second_son"})
        response = client.patch("/api/heirs/second_son", json={"co_resident": True})

    assert response.status_code == 200
    body = response.json()
    assert body["state"]["home_acquirer_id"] == "second_son"
    assert body["analysis"]["home_acquirer"]["co_resident"] is True
    assert body["analysis"]["acquirer"]["id"] == "co_resident"
    assert body["analysis"]["eligibility_alerts"] == []


def test_heir_registration_adds_one_card_at_a_time() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post("/api/demo/clear-heirs")
        first = client.post(
            "/api/heirs",
            json={"relationship": "eldest_son", "co_resident": True},
        )
        second = client.post(
            "/api/heirs",
            json={"relationship": "second_son", "co_resident": False},
        )
        selected = client.patch("/api/case", json={"home_acquirer_id": "second_son"})

    assert first.status_code == 200
    first_body = first.json()
    assert first_body["state"]["home_acquirer_id"] == "eldest_son"
    assert first_body["analysis"]["home_acquirer"]["name"] == "長男"
    assert first_body["analysis"]["acquirer"]["id"] == "co_resident"

    assert second.status_code == 200
    second_body = second.json()
    heirs = {heir["id"]: heir for heir in second_body["state"]["heirs"]}
    assert heirs["eldest_son"]["co_resident"] is True
    assert heirs["second_son"]["co_resident"] is False
    assert second_body["state"]["home_acquirer_id"] == "eldest_son"

    assert selected.status_code == 200
    selected_body = selected.json()
    assert selected_body["state"]["home_acquirer_id"] == "second_son"
    assert selected_body["analysis"]["acquirer"]["id"] == "house_lost"
    assert selected_body["analysis"]["eligibility_alerts"][0]["impact_yen"] == 56000000


def test_spouse_card_selection_surfaces_secondary_inheritance_alert() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post("/api/demo/clear-heirs")
        response = client.post(
            "/api/heirs",
            json={"relationship": "spouse", "co_resident": True},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["state"]["home_acquirer_id"] == "spouse"
    assert body["analysis"]["acquirer"]["id"] == "spouse"
    alert = body["analysis"]["secondary_inheritance_alert"]
    assert alert["title"] == "二次相続の検討はされましたか？"
    assert "税理士の判断事項" in alert["message"]
    assert body["last_run"] is None


def test_heir_registration_rejects_family_nickname_relationship() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/heirs",
            json={"relationship": "mother", "co_resident": True},
        )

    assert response.status_code == 422


def test_consultation_run_fires_actions_without_reply() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/run",
            json={
                "text": "父が亡くなり、長男が同居して自宅を相続予定です。登記が祖父名義かもしれません。"
            },
        )

    assert response.status_code == 200
    body = response.json()
    run = body["run"]
    case = body["case"]
    assert run["mode"] == "deterministic_replay"
    assert run["responds_with"] == "actions_only"
    assert run["assistant_reply"] == ""
    assert run["approval_status"] == "PENDING_APPROVAL"
    assert [step["label"] for step in run["steps"]] == ["Intake", "Router", "Evidence", "Draft", "Review"]
    assert run["steps"][-1]["status"] == "PENDING_APPROVAL"
    assert case["analysis"]["acquirer"]["id"] == "co_resident"
    assert case["approval"]["word_export_enabled"] is False
    evidence_actions = run["steps"][2]["actions"]
    assert any(action["type"] == "flag_title_anomaly" for action in evidence_actions)
    assert any(action["type"] == "list_missing_documents" and action["value"] for action in evidence_actions)
    assert "select_taker_branch" in {tool["name"] for tool in run["tool_declarations"]}


def test_consultation_run_uses_gemini_function_calling_when_key_is_set(monkeypatch) -> None:
    calls = []

    class FakeInteractions:
        def create(self, **kwargs):
            calls.append(kwargs)
            return SimpleNamespace(
                steps=[
                    SimpleNamespace(
                        type="function_call",
                        id="call_router_1",
                        name="select_taker_branch",
                        arguments={
                            "acquirer_type": "house_lost",
                            "reason": "別居で賃貸暮らしの取得予定者を検出",
                        },
                    )
                ]
            )

    class FakeClient:
        interactions = FakeInteractions()

    monkeypatch.setenv("GEMINI_API_KEY", "test-key")
    monkeypatch.setattr(agent_run, "_create_gemini_client", lambda api_key: FakeClient())

    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/run",
            json={"text": "長女は別居で賃貸暮らしです。長女が自宅を相続する予定です。"},
        )

    assert response.status_code == 200
    body = response.json()
    run = body["run"]
    assert calls
    assert calls[0]["model"] == "gemini-3.5-flash"
    assert calls[0]["tools"][0]["type"] == "function"
    assert calls[0]["tools"][0]["name"] == "select_taker_branch"
    assert run["mode"] == "gemini_function_calling"
    assert run["gemini"]["used"] is True
    assert run["gemini"]["tool_name"] == "select_taker_branch"
    assert run["gemini"]["arguments"]["acquirer_type"] == "house_lost"
    assert isinstance(run["gemini"]["latency_ms"], int)
    assert body["case"]["analysis"]["acquirer"]["id"] == "house_lost"


def test_card_review_uses_gemini_function_calling_when_key_is_set(monkeypatch) -> None:
    calls = []

    class FakeInteractions:
        def create(self, **kwargs):
            calls.append(kwargs)
            return SimpleNamespace(
                steps=[
                    SimpleNamespace(
                        type="function_call",
                        id="call_router_card",
                        name="select_taker_branch",
                        arguments={
                            "acquirer_type": "co_resident",
                            "reason": "カード内容から同居親族取得と推定",
                        },
                    )
                ]
            )

    class FakeClient:
        interactions = FakeInteractions()

    monkeypatch.setenv("GEMINI_API_KEY", "test-key")
    monkeypatch.setattr(agent_run, "_create_gemini_client", lambda api_key: FakeClient())

    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.patch("/api/case", json={"home_acquirer_id": "second_son"})
        response = client.post("/api/review/from-cards")

    assert response.status_code == 200
    body = response.json()
    run = body["run"]
    assert calls
    assert calls[0]["model"] == "gemini-3.5-flash"
    assert "相談文なし" in calls[0]["input"]
    assert "自宅取得者は次男" in calls[0]["input"]
    assert run["source"] == "heir_cards"
    assert run["mode"] == "gemini_function_calling"
    assert run["gemini"]["used"] is True
    assert body["case"]["analysis"]["home_acquirer"]["name"] == "次男"
    assert body["case"]["analysis"]["acquirer"]["id"] == "house_lost"
    assert body["case"]["analysis"]["eligibility_alerts"][0]["impact_yen"] == 56000000


def test_consultation_run_falls_back_without_gemini_key(monkeypatch) -> None:
    def fail_if_called(api_key: str):
        raise AssertionError("Gemini client should not be created without GEMINI_API_KEY")

    monkeypatch.setattr(agent_run, "_create_gemini_client", fail_if_called)

    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。"},
        )

    assert response.status_code == 200
    run = response.json()["run"]
    assert run["mode"] == "deterministic_replay"
    assert run["gemini"]["used"] is False
    assert run["gemini"]["fallback_reason"] == "gemini_api_key_not_set"
    assert response.json()["case"]["analysis"]["acquirer"]["id"] == "co_resident"


def test_consultation_run_reselects_existing_home_acquirer_from_text() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        second_son = client.post(
            "/api/run",
            json={"text": "次男が自宅を相続予定です。"},
        )
        client.post("/api/demo/seed")
        eldest_son = client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。"},
        )

    assert second_son.status_code == 200
    second_body = second_son.json()
    assert second_body["case"]["state"]["home_acquirer_id"] == "second_son"
    assert second_body["case"]["analysis"]["home_acquirer"]["name"] == "次男"
    assert second_body["case"]["analysis"]["acquirer"]["id"] == "house_lost"
    assert second_body["case"]["analysis"]["eligibility_alerts"][0]["impact_yen"] == 56000000
    router_actions = second_body["run"]["steps"][1]["actions"]
    assert router_actions[0]["value"] == "house_lost"
    assert "既存の相続人カード" in router_actions[0]["reason"]

    assert eldest_son.status_code == 200
    eldest_body = eldest_son.json()
    assert eldest_body["case"]["state"]["home_acquirer_id"] == "eldest_son"
    assert eldest_body["case"]["analysis"]["home_acquirer"]["name"] == "長男"
    assert eldest_body["case"]["analysis"]["acquirer"]["id"] == "co_resident"


def test_consultation_run_populates_heir_cards_when_unregistered() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        cleared = client.post("/api/demo/clear-heirs")
        response = client.post(
            "/api/run",
            json={"text": "父が亡くなり、母が自宅を相続予定です。他の相続人は長男と次男です。"},
        )

    assert cleared.status_code == 200
    assert cleared.json()["case"]["state"]["heirs"] == []
    assert response.status_code == 200
    body = response.json()
    state = body["case"]["state"]
    heirs = {heir["id"]: heir for heir in state["heirs"]}
    assert set(heirs) == {"mother", "eldest_son", "second_son"}
    assert state["home_acquirer_id"] == "mother"
    assert body["case"]["analysis"]["acquirer"]["id"] == "spouse"
    assert body["case"]["analysis"]["home_acquirer"]["name"] == "母"
    intake_actions = body["run"]["steps"][0]["actions"]
    populate = next(action for action in intake_actions if action["type"] == "populate_heir_cards")
    assert populate["value"] == ["母", "長男", "次男"]
    assert populate["selected_home_acquirer"] == "母"
    review_actions = body["run"]["steps"][-1]["actions"]
    assert any(action["type"] == "ask_secondary_inheritance_review" for action in review_actions)


def test_consultation_run_populates_cards_and_flags_denied_value() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post("/api/demo/clear-heirs")
        response = client.post(
            "/api/run",
            json={"text": "長男は同居しています。次男は別居で自宅を相続予定です。"},
        )

    assert response.status_code == 200
    body = response.json()
    state = body["case"]["state"]
    heirs = {heir["id"]: heir for heir in state["heirs"]}
    assert heirs["eldest_son"]["co_resident"] is True
    assert heirs["second_son"]["co_resident"] is False
    assert state["home_acquirer_id"] == "second_son"
    assert body["case"]["analysis"]["acquirer"]["id"] == "house_lost"
    alerts = body["case"]["analysis"]["eligibility_alerts"]
    assert alerts[0]["impact_yen"] == 56000000
    review_actions = body["run"]["steps"][-1]["actions"]
    alert_action = next(action for action in review_actions if action["type"] == "alert_small_residence_ineligible")
    assert alert_action["impact_label"] == "課税価格"
    assert alert_action["impact_yen"] == 56000000


def test_unregistered_card_inference_avoids_spouse_substring_false_positive() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post("/api/demo/clear-heirs")
        response = client.post(
            "/api/run",
            json={"text": "大丈夫です。長女が同居して自宅を相続予定です。"},
        )

    assert response.status_code == 200
    body = response.json()
    heirs = body["case"]["state"]["heirs"]
    assert [heir["id"] for heir in heirs] == ["eldest_daughter"]
    assert body["case"]["state"]["home_acquirer_id"] == "eldest_daughter"
    assert body["case"]["analysis"]["acquirer"]["id"] == "co_resident"


def test_consultation_run_routes_house_lost_branch() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/run",
            json={"text": "別居して賃貸暮らしの長女が自宅を相続する予定です。遺産分割は見込みです。"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["case"]["analysis"]["acquirer"]["id"] == "house_lost"
    required = body["case"]["analysis"]["acquirer"]["required_document_ids"]
    assert "house_lost_no_home_docs" in required
    assert body["case"]["state"]["partition_status"] == "expected"


def test_secondary_inheritance_prompt_only_for_spouse_review() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        spouse = client.post(
            "/api/run",
            json={"text": "配偶者が取得します。登記を確認中です。"},
        )
        client.post("/api/demo/seed")
        mother_spouse = client.post(
            "/api/run",
            json={"text": "父が亡くなり、母が自宅を相続予定です。他の相続人は長男と次男です。"},
        )
        client.post("/api/demo/seed")
        co_resident = client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記を確認中です。"},
        )
        client.post("/api/demo/seed")
        house_lost = client.post(
            "/api/run",
            json={"text": "別居して賃貸暮らしの長女が自宅を相続する予定です。"},
        )

    for response in [spouse, mother_spouse]:
        assert response.status_code == 200
        review_actions = response.json()["run"]["steps"][-1]["actions"]
        secondary = next(
            action for action in review_actions if action["type"] == "ask_secondary_inheritance_review"
        )
        assert secondary["value"] == "二次相続の検討はされましたか？"
        assert "税理士の判断事項" in secondary["why"]
        assert "5,600" not in secondary["why"]
        assert "56000000" not in secondary["why"]
        assert response.json()["case"]["analysis"]["acquirer"]["id"] == "spouse"

    for response in [co_resident, house_lost]:
        assert response.status_code == 200
        actions = response.json()["run"]["steps"][-1]["actions"]
        assert not any(action["type"] == "ask_secondary_inheritance_review" for action in actions)


def test_consultation_run_keyword_regressions() -> None:
    cases = [
        ("大丈夫です。長女が同居して自宅を相続予定です。", "co_resident", "in_progress"),
        ("確定申告の準備中です。長男が同居して自宅を相続予定です。", "co_resident", "in_progress"),
        ("分割は未確定です。長男が同居して自宅を相続予定です。", "co_resident", "in_progress"),
        ("祖父が取得した土地かもしれません。長男が同居して自宅を相続予定です。", "co_resident", "in_progress"),
        ("長男が同居して自宅を相続予定です。次男は別居で賃貸です。", "co_resident", "in_progress"),
        ("warehouse lost の資料名です。長男が同居して自宅を相続予定です。", "co_resident", "in_progress"),
    ]
    with TestClient(app) as client:
        for text, expected_acquirer, expected_partition in cases:
            client.post("/api/demo/seed")
            response = client.post("/api/run", json={"text": text})
            assert response.status_code == 200, text
            state = response.json()["case"]["state"]
            assert state["acquirer_type"] == expected_acquirer, text
            assert state["partition_status"] == expected_partition, text


def test_consultation_run_document_event_regressions() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        title_response = client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記が母名義かもしれません。"},
        )
        client.post("/api/demo/seed")
        address_response = client.post(
            "/api/run",
            json={"text": "住所を尋ねています。長男が同居して自宅を相続予定です。"},
        )

    assert title_response.status_code == 200
    title_body = title_response.json()
    assert title_body["case"]["state"]["acquirer_type"] == "co_resident"
    assert title_body["case"]["state"]["documents"]["prior_generation_title_check"] == "requested"
    assert any(
        action["type"] == "flag_title_anomaly"
        for action in title_body["run"]["steps"][2]["actions"]
    )

    assert address_response.status_code == 200
    address_documents = address_response.json()["case"]["state"]["documents"]
    assert address_documents["resident_record"] == "requested"


def test_consultation_run_rejects_short_or_blank_text_in_japanese() -> None:
    with TestClient(app) as client:
        blank = client.post("/api/run", json={"text": "        "})
        short = client.post("/api/run", json={"text": "短文"})

    assert blank.status_code == 422
    assert short.status_code == 422
    assert "相談文は8文字以上" in blank.json()["detail"][0]["msg"]
    assert "相談文は8文字以上" in short.json()["detail"][0]["msg"]


def test_consultation_run_uses_neutral_case_title_and_blocks_title_taint() -> None:
    forbidden = "名義預金です"
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        response = client.post(
            "/api/run",
            json={"text": "名義預金です。長男が同居して自宅を相続予定です。"},
        )
        approved = client.post("/api/approve")
        exported = client.get("/api/export/word")

    assert response.status_code == 200
    assert response.json()["case"]["analysis"]["case"]["title"] == "小規模宅地 要件確認案件"
    assert forbidden not in response.json()["case"]["analysis"]["case"]["title"]
    assert approved.status_code == 200
    assert exported.status_code == 200

    document = Document(BytesIO(exported.content))
    combined = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert forbidden not in combined


def test_manual_state_change_clears_stale_run_and_approval() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        run = client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記を確認中です。"},
        )
        approve = client.post("/api/approve")
        patched = client.patch("/api/case", json={"acquirer_type": "spouse"})

    assert run.status_code == 200
    assert approve.status_code == 200
    assert patched.status_code == 200
    body = patched.json()
    assert body["last_run"] is None
    assert body["approval"]["word_export_enabled"] is False
    assert body["approval"]["status"] == "NEEDS_REVIEW"


def test_word_export_requires_human_approval() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        blocked = client.get("/api/export/word")
        early_approval = client.post("/api/approve")
        client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記を確認中です。"},
        )
        approval = client.post("/api/approve")
        exported = client.get("/api/export/word")

    assert blocked.status_code == 409
    assert blocked.json()["detail"] == "approval_required"
    assert early_approval.status_code == 409
    assert early_approval.json()["detail"] == "review_not_ready"
    assert approval.status_code == 200
    assert approval.json()["approval"]["word_export_enabled"] is True
    assert exported.status_code == 200
    assert exported.content[:2] == b"PK"


def test_card_review_without_consultation_enables_approval_and_word_export() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        run = client.post("/api/review/from-cards")
        approval = client.post("/api/approve")
        exported = client.get("/api/export/word")

    assert run.status_code == 200
    body = run.json()
    assert body["run"]["source"] == "heir_cards"
    assert body["run"]["mode"] == "deterministic_replay"
    assert "相談文なし" in body["run"]["input_text"]
    assert body["run"]["steps"][-1]["summary"].startswith("税理士レビューで停止中")
    assert body["case"]["approval"]["review_ready"] is True
    assert body["case"]["approval"]["word_export_enabled"] is False
    assert approval.status_code == 200
    assert approval.json()["approval"]["word_export_enabled"] is True
    assert exported.status_code == 200
    assert exported.content[:2] == b"PK"


def test_card_review_requires_registered_heir_cards() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post("/api/demo/clear-heirs")
        response = client.post("/api/review/from-cards")

    assert response.status_code == 409
    assert response.json()["detail"] == "heirs_required_for_review"


def test_manual_overall_opinion_is_human_input_and_exports_to_word() -> None:
    manual_text = "根拠資料を確認し、別紙のとおり税理士が総合所見を記入しました。"
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記を確認中です。"},
        )
        saved = client.patch(
            "/api/manual/overall-opinion",
            json={"overall_opinion": manual_text},
        )
        approval = client.post("/api/approve")
        exported = client.get("/api/export/word")

    assert saved.status_code == 200
    saved_body = saved.json()
    assert saved_body["manual_inputs"]["overall_opinion"] == manual_text
    assert saved_body["analysis"]["draft"]["section_5_overall_opinion"] == ""
    assert saved_body["approval"]["word_export_enabled"] is False
    assert approval.status_code == 200
    assert exported.status_code == 200

    document = Document(BytesIO(exported.content))
    combined = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert manual_text in combined
    assert "（税理士が記入する欄。AIは自動記入しない。）" not in combined


def test_new_consultation_run_clears_manual_overall_opinion() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.patch(
            "/api/manual/overall-opinion",
            json={"overall_opinion": "前案件の総合所見"},
        )
        response = client.post(
            "/api/run",
            json={"text": "配偶者が取得します。登記を確認中です。"},
        )

    assert response.status_code == 200
    body = response.json()
    assert body["case"]["manual_inputs"]["overall_opinion"] == ""
    assert body["case"]["analysis"]["draft"]["section_5_overall_opinion"] == ""


def test_word_export_returns_valid_docx() -> None:
    with TestClient(app) as client:
        client.post("/api/demo/seed")
        client.post(
            "/api/run",
            json={"text": "長男が同居して自宅を相続予定です。登記を確認中です。"},
        )
        client.post("/api/approve")
        response = client.get("/api/export/word")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.content[:2] == b"PK"

    document = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    combined = text + "\n" + table_text
    assert "申告書の作成に関する計算事項等記載書面（資）" in combined
    assert "相続税申告書（令和○年○月○日相続開始分）に係る" in combined
    assert "33の2①（資）" in combined
    assert "小規模宅地等の特例" in combined
    assert "1　提示を受けた書類等に関する事項" in combined
    assert "3　計算し、整理した主な事項" in combined
    assert "区　分" in combined
    assert "事　項" in combined
    assert "備　考" in combined
    assert "4　相談に応じた事項" in combined
    assert "5　総合所見" in combined
    assert "税理士が記入する欄" in combined
    assert "否認インパクト" in combined
    assert "追徴" not in combined
