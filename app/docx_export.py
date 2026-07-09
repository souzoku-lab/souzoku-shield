from __future__ import annotations

from io import BytesIO
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Yu Mincho"
CONTENT_WIDTH_DXA = 10080
BLACK = "000000"
HEADER_FILL = "F2F2F2"
NOTE_FILL = "F8FBFF"
CAUTION_FILL = "FFF4D8"
TABLE_BORDER = "000000"


def build_shomen_docx(analysis: dict[str, Any], harness: dict[str, Any]) -> bytes:
    """書面添付の公式様式に寄せた確認中ドラフトをWord形式で生成する。"""
    doc = Document()
    _configure_document(doc)
    _add_header_footer(doc)

    _add_cover_form(doc, analysis, harness)
    _add_section_1(doc, analysis)
    _add_section_2(doc, analysis)

    doc.add_page_break()
    _add_continuation_header(doc)
    _add_section_3(doc, analysis)

    doc.add_page_break()
    _add_continuation_header(doc)
    _add_section_4(doc, analysis)
    _add_section_5(doc, analysis)
    _add_section_6(doc, analysis)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(0.62)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.24)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=9.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, before, after in [
        ("Title", 15, 0, 4),
        ("Subtitle", 10, 0, 4),
        ("Heading 1", 10, 4, 2),
        ("Heading 2", 9.5, 3, 2),
        ("Heading 3", 9.5, 2, 1),
    ]:
        style = doc.styles[style_name]
        _set_style_font(style, size=size)
        style.font.bold = style_name in {"Title", "Heading 1"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def _add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "税理士確認前 / 公開デモ用"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_font(header, size=8, color="666666")

    footer = section.footer.paragraphs[0]
    footer.text = "小規模宅地等の特例 要件確認ドラフト"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(footer, size=8, color="666666")


def _add_cover_form(doc: Document, analysis: dict[str, Any], harness: dict[str, Any]) -> None:
    _add_form_code(doc)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("相続税申告書（令和○年○月○日相続開始分）に係る")
    _set_run_font(run, size=10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("申告書の作成に関する計算事項等記載書面（資）")
    _set_run_font(run, size=15, bold=True)

    _add_case_summary_table(doc, analysis, harness)

    declaration = doc.add_paragraph()
    declaration.paragraph_format.space_before = Pt(5)
    declaration.paragraph_format.space_after = Pt(4)
    text = (
        "私（当法人）が申告書の作成に関し、計算し、整理し、又は相談に応じた事項は、"
        "下記の1から5までに掲げる事項です。"
    )
    run = declaration.add_run(text)
    _set_run_font(run, size=9.5)

    _add_notice_box(
        doc,
        "本書は公開デモ用の確認中ドラフトです。適用可否、評価額、限度面積、総合所見は税理士が確認・記入します。",
        CAUTION_FILL,
    )


def _add_form_code(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    _format_table(table, [6780, 1320, 1980], border_color="FFFFFF", cell_margin=60)
    _set_cell_text(table.cell(0, 0), "", size=8)
    _set_cell_text(table.cell(0, 1), "相続税", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(0, 2), "33の2①（資）", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_borders(table.cell(0, 1), color=TABLE_BORDER)
    _set_cell_borders(table.cell(0, 2), color=TABLE_BORDER)


def _add_case_summary_table(doc: Document, analysis: dict[str, Any], harness: dict[str, Any]) -> None:
    case = analysis["case"]
    acquirer = analysis["acquirer"]
    completion = analysis["completion"]
    land = case["land"]
    rows = [
        ("税理士又は税理士法人", "（入力欄）", "書面作成に係る税理士", "（入力欄）"),
        ("所属税理士会等", "（入力欄）", "税務代理権限証書の提出", "有 ・ 無"),
        ("依頼者", "（入力欄）", "相続税の場合（被相続人）", "（入力欄）"),
        ("案件ID", str(case["id"]), "案件名", str(case["title"])),
        (
            "土地",
            f"{land['name']} / {land['area_sqm']}㎡ / 架空評価 {land['estimated_value_yen']:,}円",
            "取得者区分",
            f"{acquirer['label']}（{acquirer['summary']}）",
        ),
        (
            "確認状況",
            f"{completion['percent']}%（{completion['label']}）",
            "否認インパクト",
            "全検査OK"
            if harness["ok"]
            else f"要確認: 課税価格 +{harness['total_damage_yen']:,}円（失われる評価減）",
        ),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    _format_table(table, [1600, 3540, 2100, 2840], cell_margin=80)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            is_label = col_idx in {0, 2}
            _set_cell_text(cell, value, size=8.5, bold=is_label)
            if is_label:
                _shade_cell(cell, HEADER_FILL)
    doc.add_paragraph()


def _add_section_1(doc: Document, analysis: dict[str, Any]) -> None:
    presented = analysis["draft"]["section_1_presented_documents"]
    table = doc.add_table(rows=4, cols=2)
    _format_table(table, [5040, 5040], cell_margin=90)
    title = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_text(title, "1　提示を受けた書類等に関する事項", size=9.5, bold=True)
    _shade_cell(title, HEADER_FILL)
    _set_cell_text(table.cell(1, 0), "書類等の名称", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(1, 1), "左記の書類等以外の書類等", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_lines_to_cell(table.cell(2, 0), presented, size=8.5, bullet=True)
    _add_lines_to_cell(
        table.cell(2, 1),
        [
            "別添チェックシートの確認書類に同じ。",
            "未提示資料は「不足資料・次の一手」で収集状況を管理する。",
        ],
        size=8.5,
        bullet=True,
    )
    note = table.cell(3, 0).merge(table.cell(3, 1))
    _set_cell_text(
        note,
        "（注）提示資料名は、申告書の作成に関し、計算し、又は整理するために用いたものに限る。",
        size=8,
    )
    doc.add_paragraph()


def _add_section_2(doc: Document, analysis: dict[str, Any]) -> None:
    prepared = analysis["draft"]["section_2_prepared_documents"]
    rows = max(len(prepared), 1) + 2
    table = doc.add_table(rows=rows, cols=2)
    _format_table(table, [4200, 5880], cell_margin=90)
    title = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_text(title, "2　自ら作成記入した書類等に関する事項", size=9.5, bold=True)
    _shade_cell(title, HEADER_FILL)
    _set_cell_text(table.cell(1, 0), "書類等の名称", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(1, 1), "作成記入の基礎となった書類等", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for index, item in enumerate(prepared or ["小規模宅地等の特例 要件確認表"], start=2):
        _set_cell_text(table.cell(index, 0), item, size=8.5)
        _set_cell_text(
            table.cell(index, 1),
            "提示資料、取得者区分、遺産分割協議の進捗、資料ステータスを基礎として確認中。",
            size=8.5,
        )
    doc.add_paragraph()


def _add_continuation_header(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    _format_table(table, [7200, 1200, 1680], border_color="FFFFFF", cell_margin=50)
    _set_cell_text(table.cell(0, 0), "", size=8)
    _set_cell_text(table.cell(0, 1), "※整理番号", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(0, 2), "", size=8)
    _set_cell_borders(table.cell(0, 1), color=TABLE_BORDER)
    _set_cell_borders(table.cell(0, 2), color=TABLE_BORDER)


def _add_section_3(doc: Document, analysis: dict[str, Any]) -> None:
    land_lines = _limit_lines(analysis["draft"]["section_3_land_review"], max_items=10)
    evidence_lines = _document_evidence_lines(analysis)

    table = doc.add_table(rows=4, cols=4)
    _format_table(table, [520, 1700, 5280, 2580], cell_margin=90)
    title = table.cell(0, 0).merge(table.cell(0, 3))
    _set_cell_text(title, "3　計算し、整理した主な事項", size=9.5, bold=True)
    _shade_cell(title, HEADER_FILL)

    _set_cell_text(table.cell(1, 0), "", size=8.5)
    _set_cell_text(table.cell(1, 1), "区　分", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(1, 2), "事　項", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(1, 3), "備　考", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_cell_text(table.cell(2, 0), "（1）", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(2, 1), "土地", size=8.5)
    _add_lines_to_cell(table.cell(2, 2), land_lines, size=8.5, bullet=True)
    _add_lines_to_cell(table.cell(2, 3), evidence_lines, size=8.2)

    _set_cell_text(table.cell(3, 0), "（2）", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(3, 1), "個別的・特徴的な事項", size=8.5)
    _add_lines_to_cell(
        table.cell(3, 2),
        [
            "取得者区分を変更した場合の必要資料差分は、反実仮想ブランチで確認中。",
            "不足資料が残る場合は、結論ではなく確認事項として記載する。",
        ],
        size=8.5,
        bullet=True,
    )
    _add_lines_to_cell(
        table.cell(3, 3),
        [
            "取得者別要件表",
            "資料収集カンバン",
            "否認インパクトハーネス",
        ],
        size=8.2,
    )
    _add_notice_box(
        doc,
        "この欄は、小規模宅地等の特例に関する計算・整理事項の候補です。断定表現を避け、税理士確認前の論点整理として出力します。",
        NOTE_FILL,
    )


def _add_section_4(doc: Document, analysis: dict[str, Any]) -> None:
    table = doc.add_table(rows=4, cols=2)
    _format_table(table, [3600, 6480], cell_margin=90)
    title = table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_text(title, "4　相談に応じた事項", size=9.5, bold=True)
    _shade_cell(title, HEADER_FILL)
    _set_cell_text(table.cell(1, 0), "事　項", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.cell(1, 1), "相　談　の　要　旨", size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_cell_text(table.cell(2, 0), "小規模宅地等の特例", size=8.5)
    _add_lines_to_cell(
        table.cell(2, 1),
        [
            f"取得者区分「{analysis['acquirer']['label']}」を前提に、必要資料と確認事項を整理中。",
            "適用可否は税理士が最終確認する旨を相談者へ説明する。",
        ],
        size=8.5,
        bullet=True,
    )
    _set_cell_text(table.cell(3, 0), "遺産分割協議", size=8.5)
    _add_lines_to_cell(
        table.cell(3, 1),
        [
            _partition_note(analysis),
            "申告期限までの資料整備状況を踏まえて、記載内容を更新する。",
        ],
        size=8.5,
        bullet=True,
    )
    doc.add_paragraph()


def _add_section_5(doc: Document, analysis: dict[str, Any]) -> None:
    manual_opinion = _manual_overall_opinion(analysis)
    lines = (
        _split_manual_lines(manual_opinion)
        if manual_opinion
        else [
            "（税理士が記入する欄。AIは自動記入しない。）",
            "根拠資料、相談経過、計算・整理事項を確認した上で、税理士の判断として記載する。",
        ]
    )
    table = doc.add_table(rows=3, cols=1)
    _format_table(table, [CONTENT_WIDTH_DXA], cell_margin=90)
    _set_cell_text(table.cell(0, 0), "5　総合所見", size=9.5, bold=True)
    _shade_cell(table.cell(0, 0), HEADER_FILL)
    _add_lines_to_cell(
        table.cell(1, 0),
        lines,
        size=8.5,
    )
    _set_row_min_height(table.rows[1], 1800)
    _set_cell_text(table.cell(2, 0), analysis["responsibility_boundary"], size=8.2)
    doc.add_paragraph()


def _manual_overall_opinion(analysis: dict[str, Any]) -> str:
    manual_inputs = analysis.get("manual_inputs", {})
    if not isinstance(manual_inputs, dict):
        return ""
    return str(manual_inputs.get("overall_opinion", "")).strip()


def _split_manual_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()] or [text.strip()]


def _add_section_6(doc: Document, analysis: dict[str, Any]) -> None:
    missing = analysis["missing_documents"]
    next_actions = analysis["next_actions"]
    rows = max(len(missing), len(next_actions), 1) + 2
    table = doc.add_table(rows=rows, cols=3)
    _format_table(table, [3000, 2100, 4980], cell_margin=90)
    title = table.cell(0, 0).merge(table.cell(0, 2))
    _set_cell_text(title, "6　その他（不足資料・次の一手）", size=9.5, bold=True)
    _shade_cell(title, HEADER_FILL)
    headers = ["不足資料", "現在の状態", "次の一手"]
    for index, header in enumerate(headers):
        _set_cell_text(table.cell(1, index), header, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for index in range(rows - 2):
        doc_item = missing[index] if index < len(missing) else {}
        action = next_actions[index] if index < len(next_actions) else ""
        _set_cell_text(table.cell(index + 2, 0), doc_item.get("label", ""), size=8.3)
        _set_cell_text(table.cell(index + 2, 1), _status_label(doc_item.get("status", "")), size=8.3)
        _set_cell_text(table.cell(index + 2, 2), action, size=8.3)


def _add_notice_box(doc: Document, text: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    _format_table(table, [CONTENT_WIDTH_DXA], border_color=TABLE_BORDER, cell_margin=90)
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    _set_cell_text(cell, text, size=8.5, bold=True)
    doc.add_paragraph()


def _document_evidence_lines(analysis: dict[str, Any]) -> list[str]:
    lines = []
    for doc_item in analysis["documents"]:
        status = _status_label(doc_item["status"])
        lines.append(f"{doc_item['label']}（{status}）")
    if not lines:
        lines.append("提示資料は未収集です。")
    return lines


def _partition_note(analysis: dict[str, Any]) -> str:
    status = analysis["state"]["partition_status"]
    return {
        "not_started": "遺産分割協議は未着手。分割見込みと申告上の取扱いを確認する。",
        "in_progress": "遺産分割協議は進行中。確定予定日又は分割見込書の要否を確認する。",
        "expected": "分割見込書を前提に、申告後の確定手続と管理予定を確認する。",
        "completed": "遺産分割協議書の内容と取得者・土地の対応関係を確認する。",
    }.get(status, "遺産分割協議の進捗を確認する。")


def _limit_lines(lines: Iterable[str], *, max_items: int) -> list[str]:
    output = [str(line) for line in lines if str(line).strip()]
    if len(output) <= max_items:
        return output
    remaining = len(output) - max_items
    return output[:max_items] + [f"ほか{remaining}件は資料収集状況に応じて追記する。"]


def _format_table(
    table: Any,
    widths: list[int],
    *,
    border_color: str = TABLE_BORDER,
    cell_margin: int = 100,
) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), border_color)

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side in ["top", "bottom", "start", "end"]:
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(cell_margin))
        node.set(qn("w:type"), "dxa")

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Inches(width / 1440)
            _set_cell_width(cell, width)


def _set_cell_text(
    cell: Any,
    text: str,
    *,
    size: float,
    bold: bool = False,
    color: str = BLACK,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, color=color)


def _add_lines_to_cell(
    cell: Any,
    lines: Iterable[str],
    *,
    size: float,
    bullet: bool = False,
    color: str = BLACK,
) -> None:
    cell.text = ""
    normalized = [str(line) for line in lines if str(line).strip()]
    if not normalized:
        normalized = [""]
    for index, line in enumerate(normalized):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.18
        text = f"・　{line}" if bullet else line
        run = paragraph.add_run(text)
        _set_run_font(run, size=size, color=color)


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_row_min_height(row: Any, height_dxa: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_dxa))
    tr_height.set(qn("w:hRule"), "atLeast")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell: Any, *, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ["top", "left", "bottom", "right"]:
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), color)


def _set_style_font(style: Any, *, size: float) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(BLACK)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia"]:
        r_fonts.set(qn(f"w:{key}"), FONT_NAME)


def _set_paragraph_font(
    paragraph: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = BLACK,
) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold, color=color)


def _set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = BLACK,
) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _status_label(status: str) -> str:
    return {
        "not_requested": "未依頼",
        "requested": "依頼済",
        "received": "受領済",
        "verified": "確認済",
    }.get(status, status)
